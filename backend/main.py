from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Header
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import sys
import warnings
import sys
import os, json, tempfile, shutil, re, math
from datetime import datetime
from urllib.parse import quote
from sqlalchemy.orm import Session
from database import get_db, Solution as DBSolution
import asyncio
from groq import Groq
from dotenv import load_dotenv
import logging
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
import requests,base64
from fastapi import Request
import asyncio
import time

from pinecone import Pinecone
from langchain_pinecone import PineconeVectorStore

#from langchain_community.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings

app = FastAPI(title="RFP Solution Generator")

# Include upload routes
# try:
from upload_routes import router as upload_router
app.include_router(upload_router)
# Include tenders routes
try:
    from tenders_routes import router as tenders_router
    app.include_router(tenders_router)
except Exception as e:
    safe_print(f"[WARN] Tenders routes not loaded: {e}")

# Include wishlist routes
try:
    from wishlist_routes import router as wishlist_router
    app.include_router(wishlist_router)
except Exception as e:
    safe_print(f"[WARN] Wishlist routes not loaded: {e}")

# Include SharePoint routes
try:
    from sharepoint_routes import router as sharepoint_router
    app.include_router(sharepoint_router)
except Exception as e:
    safe_print(f"[WARN] SharePoint routes not loaded: {e}")

# Ensure DB tables for tenders and wishlists exist on startup
try:
    from database import ensure_tenders_table, ensure_wishlists_table

    @app.on_event("startup")
    async def _startup_tables():
        try:
            ensure_tenders_table()
            ensure_wishlists_table()
        except Exception as e:
            safe_print(f"[WARN] ensure_tables failed: {e}")
except Exception as e:
    safe_print(f"[WARN] Could not register startup hook for tables: {e}")
# except Exception:
#     # If import fails during static analysis, skip. Runtime should work when packages installed.
#     pass

load_dotenv()

# Basic logging configuration (only if not configured by host)
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
if not logging.getLogger().handlers:
    logging.basicConfig(
        level=getattr(logging, LOG_LEVEL, logging.INFO),
        format="%(asctime)s %(levelname)s %(name)s - %(message)s"
    )
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("urllib3").setLevel(logging.WARNING)

# Suppress noisy third-party deprecation warnings from langchain_pinecone
warnings.filterwarnings("ignore", category=DeprecationWarning, module="langchain_pinecone")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable is required")

# --- NEW PINECONE CONFIGURATION ---
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")

if not all([PINECONE_API_KEY, PINECONE_ENVIRONMENT, PINECONE_INDEX_NAME]):
    raise ValueError("Pinecone environment variables are required")

EMBEDDING_MODEL = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# Initialize Pinecone client and vector store
pc = Pinecone(api_key=PINECONE_API_KEY)

# Use existing index
try:
    VECTOR_STORE = PineconeVectorStore(
        index_name=PINECONE_INDEX_NAME,
        embedding=EMBEDDING_MODEL
    )
except Exception as e:
    raise RuntimeError(f"Failed to connect to Pinecone index: {e}")
# --- END NEW PINECONE CONFIGURATION ---

# New: environment-driven configuration
GROQ_MODEL = os.getenv("GROQ_MODEL", "moonshotai/kimi-k2-instruct")
ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(",") if o.strip()]
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "10"))
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
AIONOS_COMPACT_OUTPUT = (os.getenv("AIONOS_COMPACT_OUTPUT", "true").lower() in ("1","true","yes"))

app.add_middleware(
    CORSMiddleware,
    allow_origins = ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

groq_client = Groq(api_key=GROQ_API_KEY, max_retries=0)

# --- SharePoint Auto-Sync Configuration ---
SHAREPOINT_AUTO_SYNC_ENABLED = (os.getenv("SHAREPOINT_AUTO_SYNC_ENABLED", "true").lower() in ("1","true","yes"))
SHAREPOINT_SYNC_INTERVAL_MINUTES = int(os.getenv("SHAREPOINT_SYNC_INTERVAL_MINUTES", "60"))  # default hourly
SHAREPOINT_INITIAL_SYNC_ON_START = (os.getenv("SHAREPOINT_INITIAL_SYNC_ON_START", "true").lower() in ("1","true","yes"))

async def _sharepoint_sync_worker() -> None:
    """Background task to keep AIonOS Knowledge Base up-to-date automatically."""
    try:
        from sharepoint_pipeline import SharePointIngestionPipeline, run_incremental_sync
        pipeline = SharePointIngestionPipeline()

        # Perform initial sync once if requested and not yet done
        if SHAREPOINT_INITIAL_SYNC_ON_START and not pipeline.delta_link:
            safe_print("[SharePoint Sync] Running initial sync on startup...")
            try:
                result = pipeline.initial_sync()
                safe_print(f"[SharePoint Sync] Initial sync finished: files={result.get('files_processed')} chunks={result.get('chunks_created')} vectors={result.get('vectors_uploaded')}")
            except Exception as e:
                safe_print(f"[SharePoint Sync] Initial sync failed: {e}")

        # Periodic incremental sync loop
        interval_seconds = max(5, SHAREPOINT_SYNC_INTERVAL_MINUTES * 60)
        while True:
            try:
                safe_print("[SharePoint Sync] Running incremental sync...")
                result = run_incremental_sync()
                # result is a dict with stats
                safe_print(f"[SharePoint Sync] Incremental sync completed: files_processed={result.get('files_processed')} files_updated={result.get('files_updated')} files_deleted={result.get('files_deleted')} vectors={result.get('vectors_uploaded')}")
            except Exception as e:
                safe_print(f"[SharePoint Sync] Incremental sync failed: {e}")
            await asyncio.sleep(interval_seconds)
    except Exception as e:
        safe_print(f"[SharePoint Sync] Worker crashed: {e}")

# Start background SharePoint auto-sync on startup (non-blocking)
if SHAREPOINT_AUTO_SYNC_ENABLED:
    @app.on_event("startup")
    async def _startup_sharepoint_sync():
        try:
            safe_print(f"[SharePoint Sync] Auto-sync enabled. Interval={SHAREPOINT_SYNC_INTERVAL_MINUTES} min. Initial on start={SHAREPOINT_INITIAL_SYNC_ON_START}")
            asyncio.create_task(_sharepoint_sync_worker())
        except Exception as e:
            safe_print(f"[WARN] Failed to start SharePoint sync worker: {e}")

# --- Logging helper to avoid Windows console encoding errors ---
def _safe_to_console_text(value: object) -> str:
    try:
        return str(value)
    except Exception:
        return repr(value)

def safe_print(*args, **kwargs) -> None:
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        enc = getattr(sys.stdout, "encoding", None) or "utf-8"
        sanitized_args = []
        for a in args:
            try:
                sanitized_args.append(_safe_to_console_text(a).encode(enc, errors="replace").decode(enc, errors="replace"))
            except Exception:
                sanitized_args.append(repr(a))
        try:
            print(*sanitized_args, **kwargs)
        except Exception:
            try:
                sys.stdout.write(" ".join(sanitized_args) + "\n")
            except Exception:
                pass


def _parse_currency_value(value: Optional[str]) -> Optional[float]:
    """Convert a human-formatted currency string (e.g., ₹1,200,000) to a float."""
    if not value or not isinstance(value, str):
        return None
    cleaned = re.sub(r"[^\d.,]", "", value)
    if not cleaned:
        return None
    cleaned = cleaned.replace(",", "")
    try:
        return float(cleaned)
    except ValueError:
        return None


def _format_currency(value: float, currency_symbol: str = "₹") -> str:
    """Format a numeric value into a currency string with thousands separators."""
    if value is None:
        return "N/A"
    if value.is_integer():
        amount = f"{int(value):,}"
    else:
        amount = f"{value:,.2f}"
    return f"{currency_symbol}{amount}"


def _format_list_markers(text: str) -> str:
    """Ensure numbered and bulleted lists render cleanly with one item per line."""
    if not text:
        return text
    formatted = text
    # Insert newline between sentences and numbered list markers (e.g., "... 2. Item")
    formatted = re.sub(r'(?<=\S)\s+(?=\d+\.\s)', '\n', formatted)
    # Insert newline before dash bullets appearing mid-line
    formatted = re.sub(r'(?<=\S)\s+(?=-\s)', '\n', formatted)
    # Collapse excessive blank lines
    formatted = re.sub(r'\n{3,}', '\n\n', formatted)
    return formatted.strip()

class SolutionStep(BaseModel):
    title:str
    description:str

class Milestone(BaseModel):
    phase:str
    duration:str
    description:str

class ResourceItem(BaseModel):
    role: str
    count: int
    years_of_experience: Optional[int] = None
    responsibilities: Optional[str] = None

class CostItem(BaseModel):
    item: str
    cost: str
    notes: Optional[str] = None

class KPIItem(BaseModel):
    metric: str
    target: str
    measurement_method: str
    frequency: Optional[str] = None

class GeneratedSolution(BaseModel):
    title: str
    date: str
    problem_statement: str
    key_challenges: List[str]
    solution_approach: List[SolutionStep]
    architecture_diagram: Optional[str] = None
    architecture_diagram_image: Optional[str] = None
    milestones: List[Milestone]
    technical_stack: List[str]
    objectives: List[str]
    acceptance_criteria: List[str]
    resources: List[ResourceItem]
    cost_analysis: List[CostItem]
    key_performance_indicators: List[KPIItem]

class GenerateTextBody(BaseModel):
    text: str
    method: Optional[str] = "knowledgeBase"  # 'knowledgeBase' or 'llmOnly'
    knowledge_base: Optional[str] = None  # 'AIonOS' for SharePoint, None for uploaded solutions

class ProductRecommendation(BaseModel):
    name: str
    description: str
    url: str
    score: float
 
class RecommendBody(BaseModel):
    text: str
 
class RetrievalInfo(BaseModel):
    knowledge_base: Optional[str] = None
    top_k: Optional[int] = None
    retrieved_count: int = 0
    filenames: List[str] = []

class SolutionWithRecommendations(BaseModel):
    solution: GeneratedSolution
    recommendations: List[ProductRecommendation] = []
    retrieval_info: Optional[RetrievalInfo] = None

# Document extraction functions
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading Word document: {str(e)}")
    return text

# Legacy .doc is intentionally not supported per requirements

def _get_logo_path() -> str:
    """Resolve absolute path to AIONOS_logo.png at project root."""
    backend_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(backend_dir)
    logo_path = os.path.join(project_root, 'AIONOS_logo.png')
    return logo_path

# --- AionOS Product Catalog & Lightweight Matcher ---
_AIONOS_PRODUCTS = [
    {
        "name": "IntelliMate\u2122",
        "description": (
            "IntelliMate\u2122 unifies NLP, computer vision, generative AI, predictive analytics into a single enterprise platform to orchestrate autonomous AI agents, integrate with legacy systems, improve operations and cybersecurity."
        ),
        "url": "https://aionos.io/intellimate",
        "keywords": [
            "unified ai platform","integration","legacy systems","agents","decision-making","efficiency","cybersecurity","enterprise","observability","operations"
        ],
    },
    {
        "name": "IntelliConverse",
        "description": (
            "Voice AI platform for multi-lingual conversational experiences across channels; automates customer interactions by connecting dialogues to backend systems and understanding complex intents."
        ),
        "url": "https://aionos.io/products/intelliconverse",
        "keywords": [
            "voice ai","multilingual","multi-lingual","chatbot","contact center","customer support","omnichannel","dialog","intent","automation"
        ],
    },
    {
        "name": "IntelliReach",
        "description": (
            "AI-powered marketing suite for campaign optimization, audience segmentation, personalized interactions, and analytics-driven performance insights to improve ROI."
        ),
        "url": "https://aionos.io/products/intelliconverse",
        "keywords": [
            "marketing","campaign","roi","personalization","audience segmentation","analytics","engagement","distribution","effectiveness"
        ],
    },
    {
        "name": "IntelliWorkflow",
        "description": (
            "Automates and orchestrates repetitive business tasks across departments, reducing manual processing and errors, integrating with existing workflows for operational efficiency."
        ),
        "url": "https://aionos.io/products/intelliworkflow",
        "keywords": [
            "workflow","automation","rpa","orchestration","repetitive tasks","manual","errors","process","efficiency","scale"
        ],
    },
    {
        "name": "IntelliResilience",
        "description": (
            "Autonomous AI-driven platform for business continuity, risk management, and rapid recovery with predictive risk assessment, real-time monitoring, and adaptive recovery strategies."
        ),
        "url": "https://aionos.io/products/intelliresilience",
        "keywords": [
            "resilience","business continuity","disaster recovery","outage","cyberattack","downtime","risk","recovery","compliance","incident response"
        ],
    },
    {
        "name": "IntelliPulse",
        "description": (
            "Advanced feedback and survey automation platform to collect, analyze, and act on insights; intelligent surveys and analytics uncover patterns in unstructured feedback."
        ),
        "url": "https://aionos.io/products/intellivision",
        "keywords": [
            "survey","feedback","insights","completion rate","nps","employee","customer","voice of customer","questionnaire","analytics"
        ],
    },
]
 
_STOPWORDS = set([
    "the","and","of","to","a","in","for","on","is","with","by","as","at","from","or","an","be","that","this","it","are","our","we","your","their","into","across","over","after","within","without"
])
 
def _tokenize(text: str) -> List[str]:
    cleaned = ''.join([c.lower() if (c.isalnum() or c.isspace()) else ' ' for c in (text or "")])
    return [t for t in cleaned.split() if t and t not in _STOPWORDS]
 
def _similarity(a: str, b: str) -> float:
    # Jaccard similarity over token sets as a lightweight proxy for semantic match
    a_set = set(_tokenize(a))
    b_set = set(_tokenize(b))
    if not a_set or not b_set:
        return 0.0
    inter = len(a_set & b_set)
    union = len(a_set | b_set)
    return inter / union if union else 0.0
 
def find_product_recommendations(problem_statement: str, threshold: float = 0.30) -> List[ProductRecommendation]:
    recs: List[ProductRecommendation] = []
    for p in _AIONOS_PRODUCTS:
        # base similarity on full description
        base_score = _similarity(problem_statement, p["description"]) if problem_statement else 0.0
        # lightweight keyword boost
        boost = 0.0
        kw_matches =0
        if problem_statement and p.get("keywords"):
            ps_lower = (problem_statement or "").lower()
            kw_matches = sum (1 for kw in p["keywords"] if kw in ps_lower)
            if kw_matches:
                # Normalize by keyword count; cap boost weight to avoid overpowering base score
                boost = min(1.0, kw_matches / max(3, len(p["keywords"])) ) * 0.8  # up to +0.8
        score = min(1.0, 0.6 * base_score + boost)
        # Guarantee a minimal score if at least one keyword matches
        if kw_matches >= 1:
            score = max(score, 0.35)
        if score >= threshold:
            recs.append(ProductRecommendation(name=p["name"], description=p["description"], url=p["url"], score=round(float(score), 2)))
    recs.sort(key=lambda r: r.score, reverse=True)
    return recs
# Utility: add Page X of Y footer using field codes

def _add_page_number_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Helper to append a field
    def _add_field(run, fld_char_type):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), fld_char_type)
        run._r.append(fld)

    # Build: "Page " + PAGE + " of " + NUMPAGES
    r1 = paragraph.add_run("Page ")

    r_page_begin = paragraph.add_run()
    _add_field(r_page_begin, 'begin')
    instr_page = OxmlElement('w:instrText')
    instr_page.set(qn('xml:space'), 'preserve')
    instr_page.text = ' PAGE '
    r_page_begin._r.append(instr_page)
    r_page_sep = paragraph.add_run()
    _add_field(r_page_sep, 'separate')
    r_page_text = paragraph.add_run()
    _add_field(r_page_text, 'end')

    r2 = paragraph.add_run(" of ")

    r_nump_begin = paragraph.add_run()
    _add_field(r_nump_begin, 'begin')
    instr_nump = OxmlElement('w:instrText')
    instr_nump.set(qn('xml:space'), 'preserve')
    instr_nump.text = ' NUMPAGES '
    r_nump_begin._r.append(instr_nump)
    r_nump_sep = paragraph.add_run()
    _add_field(r_nump_sep, 'separate')
    r_nump_text = paragraph.add_run()
    _add_field(r_nump_text, 'end')

# Utility: add a bookmark to a heading paragraph and add PAGEREF fields

def _add_bookmark(paragraph, name: str, bm_id: int) -> None:
    """Add a bookmark to a paragraph"""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bm_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bm_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_pageref_to_cell(cell, bookmark_name: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def _add_field(run, fld_char_type):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), fld_char_type)
        run._r.append(fld)
    r_begin = p.add_run()
    _add_field(r_begin, 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark_name} '
    r_begin._r.append(instr)
    r_sep = p.add_run()
    _add_field(r_sep, 'separate')
    r_end = p.add_run()
    _add_field(r_end, 'end')

# Utility: add a leadered index line: "NN Title ......... <PAGEREF>"
# (kept for reference but not used — we now insert a TOC field)

def _add_index_line(doc: Document, number_str: str, title: str, bookmark_name: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    tabstops = p.paragraph_format.tab_stops
    tabstops.add_tab_stop(Inches(6.0), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
    run_left = p.add_run(f"{number_str} {title}\t")
    def _add_field(run, fld_char_type):
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), fld_char_type)
        run._r.append(fld)
    r_begin = p.add_run()
    _add_field(r_begin, 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark_name} '
    r_begin._r.append(instr)
    r_sep = p.add_run()
    _add_field(r_sep, 'separate')
    r_end = p.add_run()
    _add_field(r_end, 'end')

# Insert a Table of Contents field (Heading levels 1 only)

def _insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Create field characters within runs (valid WordprocessingML)
    def fld_char(char_type: str):
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), char_type)
        return el
    def instr_text(text: str):
        el = OxmlElement('w:instrText')
        el.set(qn('xml:space'), 'preserve')
        el.text = text
        return el
    r1 = p.add_run(); r1._r.append(fld_char('begin'))
    r2 = p.add_run(); r2._r.append(instr_text(' TOC \\o "1-1" \\h \\z \\u '))
    r3 = p.add_run(); r3._r.append(fld_char('separate'))
    r4 = p.add_run(); r4._r.append(fld_char('end'))

# --- Helper Functions for Content Expansion ---

def _len_ok(text: str, min_len: int) -> bool:
    """Check if text meets minimum length requirement."""
    return len((text or "").strip()) >= min_len

def _calculate_total_response_size(solution_data: dict) -> int:
    """Calculate total character count of all fields in solution."""
    total = 0
    total += len(solution_data.get("problem_statement", ""))
    total += sum(len(str(c)) for c in solution_data.get("key_challenges", []))
    total += sum(len(str(step.get("title", ""))) + len(str(step.get("description", ""))) 
                 for step in solution_data.get("solution_approach", []))
    total += sum(len(str(obj)) for obj in solution_data.get("objectives", []))
    total += sum(len(str(crit)) for crit in solution_data.get("acceptance_criteria", []))
    total += sum(len(str(m.get("description", ""))) for m in solution_data.get("milestones", []))
    total += sum(len(str(kpi.get("metric", ""))) + len(str(kpi.get("target", ""))) 
                 for kpi in solution_data.get("key_performance_indicators", []))
    total += sum(len(str(cost.get("item", ""))) + len(str(cost.get("notes", ""))) 
                 for cost in solution_data.get("cost_analysis", []))
    return total

def _needs_expansion(sol: dict) -> bool:
    """Check if solution needs expansion based on quality requirements."""
    # Problem statement: minimum 2,500 characters
    if not _len_ok(sol.get("problem_statement", ""), 2500):
        return True
    
    # Key challenges: 14-20 items, each minimum 1,500 characters
    challenges = sol.get("key_challenges", [])
    if len(challenges) < 14:
        return True
    if any(len((c or "").strip()) < 1500 for c in challenges):
        return True
    
    # Solution approach: ≥8 steps, each description minimum 2,000 characters
    approach = sol.get("solution_approach", [])
    if len(approach) < 8:
        return True
    if any(len((step.get("description") or "").strip()) < 2000 for step in approach):
        return True
    
    # Objectives: 12-16 items, each minimum 1,000 characters
    objectives = sol.get("objectives", [])
    if len(objectives) < 12 or len(objectives) > 16:
        return True
    if any(len((obj or "").strip()) < 1000 for obj in objectives):
        return True
    
    # Acceptance criteria: 18-24 items, each minimum 1,200 characters
    criteria = sol.get("acceptance_criteria", [])
    if len(criteria) < 18 or len(criteria) > 24:
        return True
    if any(len((crit or "").strip()) < 1200 for crit in criteria):
        return True
    
    # Milestones: 8-12 phases, each minimum 500 characters
    milestones = sol.get("milestones", [])
    if len(milestones) < 8 or len(milestones) > 12:
        return True
    if any(len((m.get("description") or "").strip()) < 500 for m in milestones):
        return True
    
    # KPIs: minimum 6 items
    kpis = sol.get("key_performance_indicators", [])
    if len(kpis) < 6:
        return True
    
    # Cost analysis: minimum 6 items
    costs = sol.get("cost_analysis", [])
    if len(costs) < 6:
        return True
    
    return False

def _extract_and_parse_json(response_text: str) -> dict:
    """Extract and parse JSON from LLM response with aggressive error handling."""
    # Try fenced JSON first
    json_start = response_text.find("```json")
    json_end = response_text.rfind("```")
    if json_start != -1 and json_end != -1 and json_end > json_start:
        json_str = response_text[json_start + len("```json"):json_end].strip()
    else:
        # Try mermaid fence
        json_start = response_text.find("```mermaid")
        if json_start != -1:
            json_end = response_text.rfind("```", json_start + 10)
            if json_end > json_start:
                json_str = response_text[json_start + len("```mermaid"):json_end].strip()
            else:
                json_str = response_text[json_start + len("```mermaid"):].strip()
        else:
            # Fallback: try to slice from first { to last }
            brace_start = response_text.find('{')
            brace_end = response_text.rfind('}')
            if brace_start == -1 or brace_end == -1:
                raise ValueError("No JSON found in response")
            json_str = response_text[brace_start:brace_end+1]
    
    # Aggressive pre-fixing for common errors
    json_str = json_str.strip()
    if not json_str.startswith('{'):
        # Try to find first {
        first_brace = json_str.find('{')
        if first_brace != -1:
            json_str = json_str[first_brace:]
    
    # Remove trailing commas before closing braces/brackets
    json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
    
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        safe_print(f"[WARN] JSON parse error: {e}")
        safe_print(f"[DEBUG] JSON string (first 500 chars): {json_str[:500]}")
        # Attempt automatic repair if library available
        try:
            from json_repair import repair_json  # type: ignore

            repaired = repair_json(json_str)
            if repaired and repaired != json_str:
                safe_print("[INFO] JSON repair applied successfully.")
                return json.loads(repaired)
        except ModuleNotFoundError:
            safe_print("[INFO] json_repair library not available; skipping repair.")
        except Exception as repair_exc:
            safe_print(f"[WARN] JSON repair failed: {repair_exc}")
        raise

def _expand_solution_json(solution_data: dict, rfp_text: str) -> dict:
    """Expand solution JSON to meet professional-grade detail requirements."""
    improvement_prompt = f"""
You are improving a technical proposal JSON to professional-grade depth.

CRITICAL REQUIREMENTS - MINIMUM CHARACTER COUNTS:
- problem_statement: MINIMUM 2500 characters; MUST be multi-paragraph prose with detailed analysis
- key_challenges: MUST be 14–20 items; each item MINIMUM 1500 characters (8–12 sentences per challenge)
- solution_approach: MUST be ≥8 steps; each step's description MINIMUM 2000 characters (10–15 sentences per step)
- objectives: MUST be 12–16 items; each item MINIMUM 1000 characters (5–8 sentences per objective)
- acceptance_criteria: MUST be 18–24 items; each item MINIMUM 1200 characters (6–10 sentences per criterion)
- milestones: MUST be 8–12 phases; each description MINIMUM 500 characters
- key_performance_indicators: MUST be ≥6 items with detailed metrics
- cost_analysis: MUST be ≥6 items with detailed breakdowns
- milestones array MUST contain objects with keys "phase", "duration", "description" (no plain strings)
- cost_analysis array MUST contain objects with keys "item", "cost", "notes" (no plain strings)

RFP Context (for depth):
{rfp_text[:6000]}

Current Solution JSON (improve this):
{json.dumps(solution_data, indent=2)[:15000]}

IMPORTANT:
- Keep the EXACT same JSON schema/structure
- Expand content to meet ALL minimum character counts above
- Add domain-specific technical details and jargon where appropriate
- Make descriptions comprehensive and professional
- Do NOT change field names or structure
- Respond ONLY with valid JSON (no fences, no explanations)
"""
    
    try:
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You improve JSON to meet professional-grade detail. Output ONLY valid JSON. No fences, no comments, no explanations."},
                {"role": "user", "content": improvement_prompt}
            ],
            temperature=0.6,
            max_tokens=10000,
        )
        
        response_text = response.choices[0].message.content or ""
        expanded = _extract_and_parse_json(response_text)
        
        # Validate all required keys are present
        required_keys = ["title", "date", "problem_statement", "key_challenges", "solution_approach", 
                        "milestones", "technical_stack", "objectives", "acceptance_criteria", 
                        "resources", "cost_analysis", "key_performance_indicators"]
        for key in required_keys:
            if key not in expanded:
                safe_print(f"[WARN] Missing key '{key}' in expanded solution, using original")
                expanded[key] = solution_data.get(key)
        
        return expanded
    except Exception as e:
        safe_print(f"[WARN] Expansion failed: {e}, returning original")
        return solution_data

# --- Helper Functions for Diagram Quality ---

def _validate_mermaid_syntax(code: str) -> bool:
    """Basic validation of Mermaid syntax."""
    if not code or not code.strip():
        return False
    code_lower = code.lower().strip()
    # Must start with graph, flowchart, or erDiagram
    if not (code_lower.startswith("graph ") or code_lower.startswith("flowchart ") or 
            code_lower.startswith("erdiagram")):
        return False
    return True

_MERMAID_ID_SAFE = re.compile(r'[^A-Za-z0-9_]')


def _sanitize_mermaid_identifier(token: str, fallback_prefix: str = "Node") -> str:
    """Normalize Mermaid identifiers to alphanumeric/underscore and avoid leading digits."""
    token = (token or "").strip()
    if not token:
        return fallback_prefix
    cleaned = _MERMAID_ID_SAFE.sub("_", token)
    cleaned = cleaned.lstrip("_")
    if not cleaned:
        cleaned = fallback_prefix
    if cleaned[0].isdigit():
        cleaned = f"{fallback_prefix}{cleaned}"
    return cleaned


def _sanitize_mermaid_code(code: str | None) -> str | None:
    """Produce safer Mermaid code by normalizing identifiers and removing fences."""
    if not code:
        return code

    text = str(code).strip()
    if text.startswith("```"):
        end_idx = text.rfind("```")
        if end_idx > 0:
            inner = text[3:end_idx].strip()
            if inner.lower().startswith(("json", "mermaid")):
                inner = inner.split("\n", 1)[1] if "\n" in inner else ""
            text = inner or text
    if text.lower().startswith("mermaid"):
        text = text[7:].strip()

    text = text.replace("\r\n", "\n")

    lines: list[str] = []
    subgraph_pattern = re.compile(r"^(\s*subgraph\s+)([^\s]+)(.*)$", re.IGNORECASE)
    node_pattern = re.compile(
        r"^(\s*)([A-Za-z0-9_-]+)(\s*)(\[[^\]]+\]|\([^\)]+\)|\(\[[^\]]+\]\)|\{\{[^}]+\}\}|<[^>]+>)(.*)$"
    )
    edge_pattern = re.compile(r"([A-Za-z0-9_-]+)(\s*[-.]+(?:\>|o\>|x\>|-\||\|>|\>|)\|?.*?\|?)([A-Za-z0-9_-]+)")

    for original_line in text.split("\n"):
        line = original_line
        match = subgraph_pattern.match(line.strip())
        if match:
            identifier = _sanitize_mermaid_identifier(match.group(2), "Subgraph")
            line = f"{match.group(1)}{identifier}{match.group(3)}"
        else:
            m = node_pattern.match(line)
            if m:
                identifier = _sanitize_mermaid_identifier(m.group(2))
                line = f"{m.group(1)}{identifier}{m.group(3)}{m.group(4)}{m.group(5)}"

        def _edge_replacer(edge_match: re.Match[str]) -> str:
            src = _sanitize_mermaid_identifier(edge_match.group(1))
            arrow = edge_match.group(2)
            dest = _sanitize_mermaid_identifier(edge_match.group(3))
            return f"{src}{arrow}{dest}"

        line = edge_pattern.sub(_edge_replacer, line)
        lines.append(line)

    sanitized = "\n".join(lines).strip()
    if not sanitized.lower().startswith(("flowchart", "graph", "erdiagram")):
        sanitized = f"flowchart TD\n{sanitized}"
    return sanitized


def _diagram_is_basic(mermaid_code: str) -> bool:
    """Check if diagram is too basic (needs improvement)."""
    if not mermaid_code or not isinstance(mermaid_code, str):
        return True
    
    code = mermaid_code.strip()
    if not _validate_mermaid_syntax(code):
        return True
    
    # Count nodes (approximate by counting '[')
    node_count = code.count('[')
    if node_count < 12:
        return True
    
    # Check for subgraphs
    code_lower = code.lower()
    has_subgraph = ('subgraph' in code_lower and 'end' in code_lower)
    if not has_subgraph:
        return True
    
    return False

def _improve_diagram_mermaid(rfp_text: str, current: str) -> str:
    """Improve a basic Mermaid diagram to professional-grade."""
    prompt = f"""
You will output ONLY valid Mermaid flowchart code for a professional system architecture diagram.

CRITICAL SYNTAX RULES (MUST FOLLOW):
- Start with: flowchart TD or graph TD
- Use subgraphs: subgraph LABEL[...] ... end (must balance)
- Node IDs: simple alphanumeric only (e.g., WebApp, UserSvc, PrimaryDB)
- NO spaces or hyphens in node IDs
- Edges: use --> or -->|Label| only
- NO comments, NO code fences, NO explanations
- Apply classDef styling at the end if needed

RFP Context (for domain-specific architecture):
{rfp_text[:4000]}

Current Diagram (fix and improve):
{current[:1500]}

REQUIREMENTS:
- Minimum 15 nodes showing realistic system components
- Use 5-7 subgraphs for layers: Client, Edge, Gateway, Services, Data, Infrastructure
- All nodes MUST be connected (no isolated nodes)
- Show actual data flow and interactions
- Use domain-appropriate technologies (AWS/Azure/GCP, databases, queues, etc.)
- Include AI/ML components if applicable (LLMs, model serving, etc.)

Output ONLY the Mermaid code, nothing else.
"""
    
    try:
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "Output only valid Mermaid flowchart code. No fences, no comments, no explanations."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=1500,
        )
        
        text = response.choices[0].message.content or ""
        
        # Strip code fences if present
        text = text.strip()
        if text.startswith("```"):
            si = text.find("```")
            ei = text.rfind("```")
            if ei > si:
                text = text[si+3:ei].replace("mermaid", "").strip()
        
        # Validate
        if _validate_mermaid_syntax(text):
            return text
        else:
            safe_print("[WARN] Improved diagram failed validation, using original")
            return current
    except Exception as e:
        safe_print(f"[WARN] Diagram improvement failed: {e}, using original")
        return current

def _normalize_solution_shapes(solution: dict) -> dict:
    """Ensure solution data conforms to expected schema for Pydantic validation."""
    if not isinstance(solution, dict):
        return {}

    normalized = dict(solution)

    # Normalize key challenges to list of strings
    challenges: List[str] = []
    for item in solution.get("key_challenges") or []:
        if isinstance(item, dict):
            text = (item.get("description") or item.get("text") or "").strip()
        else:
            text = str(item).strip()
        if text:
            challenges.append(text)
    normalized["key_challenges"] = challenges

    # Normalize solution approach to list of dicts with title & description
    approach: List[dict] = []
    for idx, step in enumerate(solution.get("solution_approach") or []):
        if isinstance(step, dict):
            title = str(step.get("title") or step.get("phase") or f"Step {idx + 1}").strip()
            description = str(step.get("description") or step.get("details") or "").strip()
        else:
            text = str(step).strip()
            title = f"Step {idx + 1}"
            description = text
        approach.append({"title": title or f"Step {idx + 1}", "description": description})
    normalized["solution_approach"] = approach

    # Normalize milestones to list of dicts with phase, duration, description
    milestones: List[dict] = []
    for idx, milestone in enumerate(solution.get("milestones") or []):
        if isinstance(milestone, dict):
            phase = str(milestone.get("phase") or milestone.get("title") or f"Phase {idx + 1}").strip()
            duration = str(milestone.get("duration") or milestone.get("timeline") or "").strip()
            description = str(milestone.get("description") or milestone.get("details") or "").strip()
        else:
            text = str(milestone).strip()
            phase = text[:80] or f"Phase {idx + 1}"
            duration = ""
            description = text
        milestones.append({
            "phase": phase or f"Phase {idx + 1}",
            "duration": duration,
            "description": description
        })
    normalized["milestones"] = milestones

    # Normalize cost analysis to list of dicts with item, cost, notes
    cost_items: List[dict] = []
    for idx, cost in enumerate(solution.get("cost_analysis") or []):
        if isinstance(cost, dict):
            item_name = str(cost.get("item") or cost.get("title") or cost.get("name") or f"Cost Item {idx + 1}").strip()
            amount = str(cost.get("cost") or cost.get("amount") or "").strip()
            notes = str(cost.get("notes") or cost.get("description") or "").strip()
        else:
            text = str(cost).strip()
            item_name = text or f"Cost Item {idx + 1}"
            amount = ""
            notes = ""
        cost_items.append({
            "item": item_name or f"Cost Item {idx + 1}",
            "cost": amount,
            "notes": notes
        })
    normalized["cost_analysis"] = cost_items

    # Normalize objectives and acceptance criteria to list of strings
    normalized["objectives"] = [str(obj).strip() for obj in (solution.get("objectives") or []) if str(obj).strip()]
    normalized["acceptance_criteria"] = [str(crit).strip() for crit in (solution.get("acceptance_criteria") or []) if str(crit).strip()]

    # Normalize technical stack to strings
    normalized["technical_stack"] = [str(tech).strip() for tech in (solution.get("technical_stack") or []) if str(tech).strip()]

    # Normalize KPIs to dicts
    kpis: List[dict] = []
    for idx, kpi in enumerate(solution.get("key_performance_indicators") or []):
        if isinstance(kpi, dict):
            metric = str(kpi.get("metric") or f"KPI {idx + 1}").strip()
            target = str(kpi.get("target") or "").strip()
            measurement = str(kpi.get("measurement_method") or kpi.get("method") or "").strip()
            frequency = str(kpi.get("frequency") or kpi.get("cadence") or "").strip()
        else:
            text = str(kpi).strip()
            metric = text or f"KPI {idx + 1}"
            target = ""
            measurement = ""
            frequency = ""
        kpis.append({
            "metric": metric or f"KPI {idx + 1}",
            "target": target,
            "measurement_method": measurement,
            "frequency": frequency or None
        })
    normalized["key_performance_indicators"] = kpis

    # Normalize resources to dicts
    resources: List[dict] = []
    for idx, res in enumerate(solution.get("resources") or []):
        if isinstance(res, dict):
            role = str(res.get("role") or res.get("title") or f"Role {idx + 1}").strip()
            count_raw = res.get("count")
            try:
                count = int(count_raw)
            except (TypeError, ValueError):
                count = 0
            experience_raw = res.get("years_of_experience")
            try:
                experience = int(experience_raw) if experience_raw not in (None, "") else None
            except (TypeError, ValueError):
                experience = None
            responsibilities = str(res.get("responsibilities") or res.get("description") or "").strip()
        else:
            text = str(res).strip()
            role = text or f"Role {idx + 1}"
            count = 0
            experience = None
            responsibilities = text
        resources.append({
            "role": role or f"Role {idx + 1}",
            "count": count,
            "years_of_experience": experience,
            "responsibilities": responsibilities
        })
    normalized["resources"] = resources

    return normalized

def render_mermaid_to_image(mermaid_code: str) -> str | None:
    """Render Mermaid code to a temporary PNG file using mermaid.ink API with QuickChart fallback."""
    if not mermaid_code or not mermaid_code.strip():
        return None
    
    # Clean code
    code = _sanitize_mermaid_code(mermaid_code) or ""
    code = code.strip()
    if code.startswith('```'):
        si = code.find('```')
        ei = code.rfind('```')
        if ei > si:
            code = code[si+3:ei].replace('mermaid', '').strip()
    
    code = code.strip()
    if not code:
        return None
    
    temp_dir = tempfile.gettempdir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Primary method: mermaid.ink API
    try:
        # Base64 URL-safe encode
        encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('utf-8').rstrip('=')
        url = f"https://mermaid.ink/img/{encoded}"
        
        response = requests.get(url, timeout=15)
        if response.status_code == 200 and response.content[:4] == b'\x89PNG':
            image_path = os.path.join(temp_dir, f'architecture_diagram_{timestamp}.png')
            with open(image_path, "wb") as f:
                f.write(response.content)
            return image_path
    except Exception as e:
        safe_print(f"[WARN] mermaid.ink API failed: {e}, trying QuickChart fallback")
    
    # Fallback: QuickChart API
    try:
        encoded_code = quote(code)
        url = f"https://quickchart.io/mermaid?c={encoded_code}"
        
        response = requests.get(url, timeout=15)
        if response.status_code == 200 and response.content[:4] == b'\x89PNG':
            image_path = os.path.join(temp_dir, f'architecture_diagram_{timestamp}.png')
            with open(image_path, "wb") as f:
                f.write(response.content)
            return image_path
    except Exception as e:
        safe_print(f"[WARN] QuickChart API also failed: {e}")
    
    return None

# --- Async LLM Completion Helper ---
_RETRYABLE_STATUS_CODES = {429, 500, 502, 503, 504}
_LLM_RETRY_DELAY_SECONDS = 3
_LLM_MAX_RETRIES = 8


def _should_retry_exception(exc: Exception) -> bool:
    status = getattr(exc, "status_code", None) or getattr(exc, "http_status", None)
    if status in _RETRYABLE_STATUS_CODES:
        return True
    message = str(exc).lower()
    if "429" in message or "rate limit" in message:
        return True
    if any(term in message for term in ["temporarily unavailable", "timeout", "overload"]):
        return True
    return False


async def async_llm_complete(messages: List[dict], temperature: float = 0.3, max_tokens: int = 8000) -> str:
    """Async wrapper for Groq LLM completion with custom retry backoff."""

    def _call_with_retries() -> str:
        last_error: Exception | None = None
        for attempt in range(1, _LLM_MAX_RETRIES + 1):
            try:
                response = groq_client.chat.completions.create(
                    model=GROQ_MODEL,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens
                )
                return response.choices[0].message.content or ""
            except Exception as exc:  # noqa: BLE001
                last_error = exc
                if attempt >= _LLM_MAX_RETRIES or not _should_retry_exception(exc):
                    raise
                safe_print(
                    f"[WARN] LLM request failed (attempt {attempt}/{_LLM_MAX_RETRIES}); "
                    f"retrying in {_LLM_RETRY_DELAY_SECONDS} seconds..."
                )
                time.sleep(_LLM_RETRY_DELAY_SECONDS)
        if last_error:
            raise last_error
        raise RuntimeError("LLM completion failed without raising an explicit exception")

    return await asyncio.to_thread(_call_with_retries)

# LLM Processing
async def analyze_rfp_with_groq(rfp_text: str, use_rag: bool = True, knowledge_base: Optional[str] = None):
    """Analyze RFP text using Groq and generate solution with multi-stage expansion
    
    Args:
        rfp_text: Input RFP text or problem statement
        use_rag: Whether to use RAG retrieval
        knowledge_base: Optional knowledge base filter ('AIonOS' for SharePoint, None for uploaded solutions)
    """

    # Step 1: Retrieve relevant documents from vector store
    retrieved_docs = []
    if use_rag:
        try:
            # If knowledge_base is specified, use Pinecone client directly for metadata filtering
            if knowledge_base == "AIonOS":
                try:
                    # Use Pinecone query with metadata filter
                    query_vector = EMBEDDING_MODEL.embed_query(rfp_text)
                    results = pc.Index(PINECONE_INDEX_NAME).query(
                        vector=query_vector,
                        top_k=5,
                        include_metadata=True,
                        filter={"knowledge_base": {"$eq": "AIonOS"}}
                    )
                    # Convert to LangChain Document format
                    from langchain.schema import Document
                    retrieved_docs = [
                        Document(
                            page_content=match['metadata'].get('text', ''),
                            metadata=match['metadata']
                        )
                        for match in results['matches']
                    ]
                    safe_print(f"Retrieved {len(retrieved_docs)} documents from AIonOS knowledge base")
                except Exception as e:
                    safe_print(f"Error querying Pinecone with filter: {e}, falling back to standard search")
                    retrieved_docs = VECTOR_STORE.similarity_search(rfp_text, k=5)
            else:
                # Standard RAG without filter (uses uploaded solutions)
                retrieved_docs = VECTOR_STORE.similarity_search(rfp_text, k=5)
        except Exception as e:
            safe_print(f"Error retrieving from vector store: {str(e)}")
            retrieved_docs = []

    safe_print("--- Retrieved Chunks for Validation ---")
    if not retrieved_docs:
        safe_print("No relevant chunks found in the vector store.")
    for i, doc in enumerate(retrieved_docs):
        safe_print(f"Chunk {i+1}:")
        safe_print(doc.page_content)
        safe_print(f"Source file: {doc.metadata.get('filename', 'N/A')}")
        safe_print("-" * 50)

    context_text = "\n\n".join([doc.page_content for doc in retrieved_docs]) if retrieved_docs else "No relevant references found."

    # Build retrieval metadata for UI
    retrieval_info: Optional[RetrievalInfo] = None
    if use_rag:
        filenames: List[str] = []
        try:
            for d in retrieved_docs:
                name = d.metadata.get('filename') or d.metadata.get('source') or d.metadata.get('file_name') or 'N/A'
                if name and name not in filenames:
                    filenames.append(name)
        except Exception:
            filenames = []
        retrieval_info = RetrievalInfo(
            knowledge_base=knowledge_base,
            top_k=5,
            retrieved_count=len(retrieved_docs),
            filenames=filenames[:10]
        )
    
    # Step 2: Build enhanced prompt with detailed architecture diagram instructions
    # Analyze domain hints for architecture customization
    domain_hints = ""
    rfp_lower = rfp_text.lower()
    if any(kw in rfp_lower for kw in ["ai", "machine learning", "ml", "llm", "neural", "model"]):
        domain_hints += "Domain: AI/ML - Include LLM inference, model serving, training pipelines, vector databases.\n"
    if any(kw in rfp_lower for kw in ["data pipeline", "etl", "data processing", "analytics"]):
        domain_hints += "Domain: Data Pipeline - Include data ingestion, transformation, storage, analytics layers.\n"
    if any(kw in rfp_lower for kw in ["healthcare", "medical", "patient", "hospital"]):
        domain_hints += "Domain: Healthcare - Include HIPAA compliance, patient data security, medical records systems.\n"
    if any(kw in rfp_lower for kw in ["e-commerce", "retail", "shopping", "cart", "payment"]):
        domain_hints += "Domain: E-commerce - Include payment gateways, inventory management, recommendation engines.\n"
    
    prompt = f"""
You are an expert technical consultant producing a compact, production-ready proposal JSON.
STRICT FORMAT AND BREVITY — follow EXACTLY. ENSURE RICHNESS PER LINE (about 18–28 words per line, 2 short sentences if helpful):

- problem_statement: EXPLAIN TECHNICALLY IN 10 LINES MAX. Use newline separators.
- key_challenges: EXACTLY 7 items; each item is ≤5 lines (each line ~18–28 words, technical).
- solution_approach: EXACTLY 7 steps; each step has "title" and "description" (≤5 lines; each line ~18–28 words, technical).
- architecture_diagram: PROFESSIONAL Mermaid (flowchart TD), 12–18 nodes, 5–7 subgraphs (Client, Edge, Gateway, Services, Data, Infrastructure), all nodes connected. NO code fences.
- milestones: EXACTLY 7 phases. Each has "phase", "duration", "description" (≤4 lines; each line ~16–24 words).
- technical_stack: 10–20 items, only relevant technologies.
- objectives: EXACTLY 7 items; each ≤5 lines (each line ~18–28 words).
- acceptance_criteria: EXACTLY 7 items; each ≤4 lines (each line ~16–24 words, measurable).
- resources: 6–10 roles with "role", "count", "years_of_experience", "responsibilities" (concise).
- cost_analysis: 6–10 items with "item", "cost" (INR string like "₹750,000"), "notes".
- key_performance_indicators: EXACTLY 10 with "metric","target","measurement_method","frequency".

CRITICAL RULES:
- Respond with ONE fenced ```json block (valid JSON).
- Keep content concise within the specified line limits.
- Use domain-appropriate details inferred from the RFP and retrieved context.
- Do not omit any keys in the schema below.

RFP Context (shortened):
{rfp_text[:6000]}

Retrieved References:
{context_text[:6000]}

SCHEMA (exact keys):
{{
    "title": "Solution title",
    "date": "{datetime.now().strftime('%B %Y')}",
  "problem_statement": "Line1\\nLine2\\n... up to 10 lines",
  "key_challenges": ["Up to 5 lines per challenge (7 items total)"],
  "solution_approach": [{{"title":"Step 1","description":"Up to 5 lines"}}, {{}} ... 7 items],
  "architecture_diagram": "flowchart TD\\nsubgraph ...",
  "milestones": [{{"phase":"Phase 1","duration":"X weeks","description":"Up to 4 lines"}} ... 7 items],
  "technical_stack": ["Tech1","Tech2", "..."],
  "objectives": ["Up to 5 lines per objective (7 items total)"],
  "acceptance_criteria": ["Up to 4 lines per criterion (7 items total)"],
  "resources": [{{"role":"Role","count":1,"years_of_experience":5,"responsibilities":"Concise"}} ...],
  "cost_analysis": [{{"item":"Item","cost":"₹120,000","notes":"Concise"}} ...],
  "key_performance_indicators": [{{"metric":"Metric","target":"Target","measurement_method":"Method","frequency":"Monthly"}} ... 10 items]
}}
"""
    
    # Step 3: Initial LLM generation
    try:
        system_message = """You are a technical proposal expert. Always respond with valid JSON inside a fenced ```json block.
Ensure all content meets the minimum character requirements specified in the prompt.
Generate professional-grade, detailed content for every section."""
        
        parse_attempts = 3
        solution_data: dict | None = None
        response_text: str = ""
        for attempt in range(1, parse_attempts + 1):
            response_text = await async_llm_complete(
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,  # Slightly higher for richer content
                max_tokens=9000,
            )
            try:
                solution_data = _extract_and_parse_json(response_text)
                break
            except Exception as parse_exc:
                safe_print(f"[WARN] Failed to parse LLM JSON (attempt {attempt}/{parse_attempts}): {parse_exc}")
                if attempt < parse_attempts:
                    await asyncio.sleep(1)
                    continue
                raise

        if solution_data is None:
            raise RuntimeError("LLM did not produce valid solution data.")

        solution_data = _normalize_solution_shapes(solution_data)
        solution_data["architecture_diagram"] = _sanitize_mermaid_code(solution_data.get("architecture_diagram"))
        
        # Step 4: Size check & expansion decision
        total_size = _calculate_total_response_size(solution_data)
        safe_print(f"[INFO] Initial response size: {total_size:,} characters")
        
        # Compact mode: skip heavy expansion passes
        if not AIONOS_COMPACT_OUTPUT:
            max_expansions = 2 if total_size < 180000 else 0
            for attempt in range(max_expansions):
                try:
                    safe_print(f"[INFO] Expansion pass {attempt + 1}/{max_expansions}")
                    solution_data = await asyncio.to_thread(_expand_solution_json, solution_data, rfp_text)
                    solution_data = _normalize_solution_shapes(solution_data)
                    solution_data["architecture_diagram"] = _sanitize_mermaid_code(solution_data.get("architecture_diagram"))
                except Exception as _:
                    break
        
        # Step 6: Final processing
        # Normalize shapes (ensure data matches Pydantic schema)
        # Apply backend limits and ensure compact completeness
        def _format_multiline(text: str, target_lines: int) -> str:
            text = (text or "").strip()
            if not text:
                return "\n".join(f"TBD line {i+1}" for i in range(target_lines))
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            if len(lines) >= target_lines:
                return "\n".join(lines[:target_lines])
            sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
            if len(sentences) >= target_lines:
                return "\n".join(sentences[:target_lines])
            words = text.split()
            if not words:
                return "\n".join(f"TBD line {i+1}" for i in range(target_lines))
            chunks: list[str] = []
            idx = 0
            for remaining in range(target_lines, 0, -1):
                words_left = len(words) - idx
                if words_left <= 0:
                    chunks.append(chunks[-1] if chunks else "TBD")
                    continue
                take = max(1, math.ceil(words_left / remaining))
                chunk_words = words[idx: idx + take]
                if not chunk_words:
                    chunk_words = words[-1:]
                chunks.append(" ".join(chunk_words))
                idx += take
            return "\n".join(chunks[:target_lines])

        def _limit_lines(text: str, max_lines: int) -> str:
            return _format_multiline(text, max_lines)

        def _ensure_list_range(items: list[str], min_count: int, max_count: int, filler_prefix: str, lines: int | None = None) -> tuple[list[str], bool]:
            out: list[str] = []
            filler_used = False
            for i in (items or []):
                if len(out) >= max_count:
                    break
                txt = (str(i) if not isinstance(i, dict) else i.get("description") or i.get("text") or i.get("value") or "")
                txt = txt if lines is None else _format_multiline(txt, lines)
                if txt.strip():
                    out.append(txt.strip())
            while len(out) < min_count:
                filler = f"{filler_prefix} insight {len(out)+1}"
                filler = _format_multiline(filler, lines) if lines else filler
                out.append(filler)
                filler_used = True
            return out[:max_count], filler_used

        def _apply_backfill(items: list[str] | None, target: int, lines: int | None, filler_prefix: str) -> list[str]:
            sanitized: list[str] = []
            for entry in items or []:
                if len(sanitized) >= target:
                    break
                text = str(entry)
                sanitized.append(_format_multiline(text, lines) if lines else text.strip())
            while len(sanitized) < target:
                filler = f"{filler_prefix} insight {len(sanitized)+1}"
                sanitized.append(_format_multiline(filler, lines) if lines else filler)
            return sanitized[:target]

        def _ensure_steps(steps: list[dict], target: int) -> list[dict]:
            out = []
            for idx, s in enumerate(steps or []):
                if len(out) >= target: break
                title = (s.get("title") or f"Step {idx+1}").strip()
                desc = _format_multiline(s.get("description") or "", 5)
                out.append({"title": title, "description": desc})
            while len(out) < target:
                i = len(out) + 1
                out.append({"title": f"Step {i}", "description": "\n".join(["TBD"] * 5)})
            return out[:target]

        def _ensure_milestones(milestones: list[dict], target: int) -> list[dict]:
            out = []
            for idx, m in enumerate(milestones or []):
                if len(out) >= target: break
                out.append({
                    "phase": (m.get("phase") or f"Phase {idx+1}").strip(),
                    "duration": (m.get("duration") or "2 weeks").strip(),
                    "description": _format_multiline(m.get("description") or "", 4)
                })
            while len(out) < target:
                i = len(out) + 1
                out.append({"phase": f"Phase {i}", "duration": "2 weeks", "description": "\n".join(["TBD"] * 4)})
            return out[:target]

        def _ensure_resources(resources: list[dict], target_min: int = 6, target_max: int = 10) -> list[dict]:
            out = []
            for r in resources or []:
                out.append({
                    "role": (r.get("role") or "Engineer").strip(),
                    "count": int(r.get("count") or 1),
                    "years_of_experience": r.get("years_of_experience") if isinstance(r.get("years_of_experience"), int) else 3,
                    "responsibilities": (r.get("responsibilities") or "TBD").strip()
                })
                if len(out) >= target_max: break
            while len(out) < target_min:
                out.append({"role": "Engineer", "count": 1, "years_of_experience": 3, "responsibilities": "TBD"})
            return out[:target_max]

        def _ensure_costs(costs: list[dict], target_min: int = 6, target_max: int = 10) -> list[dict]:
            out = []
            for c in costs or []:
                out.append({
                    "item": (c.get("item") or "Item").strip(),
                    "cost": (c.get("cost") or "₹100,000").strip(),
                    "notes": (c.get("notes") or "TBD").strip()
                })
                if len(out) >= target_max: break
            while len(out) < target_min:
                i = len(out) + 1
                out.append({"item": f"Item {i}", "cost": "₹100,000", "notes": "TBD"})
            return out[:target_max]

        def _ensure_kpis(kpis: list[dict], target: int = 10) -> list[dict]:
            out = []
            for k in kpis or []:
                out.append({
                    "metric": (k.get("metric") or f"KPI {len(out)+1}").strip(),
                    "target": (k.get("target") or "Target").strip(),
                    "measurement_method": (k.get("measurement_method") or "Monitoring").strip(),
                    "frequency": (k.get("frequency") or "Monthly").strip()
                })
                if len(out) >= target: break
            while len(out) < target:
                i = len(out) + 1
                out.append({"metric": f"KPI {i}", "target": "Target", "measurement_method": "Monitoring", "frequency": "Monthly"})
            return out[:target]

        # Enforce compact rules
        solution_data["problem_statement"] = _limit_lines(solution_data.get("problem_statement") or "", 10)
        key_challenges, key_fillers = _ensure_list_range(solution_data.get("key_challenges") or [], 7, 7, "Challenge", 5)
        solution_data["key_challenges"] = key_challenges
        solution_data["solution_approach"] = _ensure_steps(solution_data.get("solution_approach") or [], 7)
        solution_data["milestones"] = _ensure_milestones(solution_data.get("milestones") or [], 7)
        technical_stack, stack_fillers = _ensure_list_range(solution_data.get("technical_stack") or [], 10, 20, "Technology")
        solution_data["technical_stack"] = technical_stack
        objectives, obj_fillers = _ensure_list_range(solution_data.get("objectives") or [], 7, 7, "Objective", 5)
        solution_data["objectives"] = objectives
        acceptance, acc_fillers = _ensure_list_range(solution_data.get("acceptance_criteria") or [], 7, 7, "Criterion", 4)
        solution_data["acceptance_criteria"] = acceptance
        solution_data["resources"] = _ensure_resources(solution_data.get("resources") or [])
        solution_data["cost_analysis"] = _ensure_costs(solution_data.get("cost_analysis") or [])
        solution_data["key_performance_indicators"] = _ensure_kpis(solution_data.get("key_performance_indicators") or [])

        # Lightweight backfill if placeholders detected
        def _looks_placeholder_list(values: list[str], prefix: str) -> bool:
            sample = " ".join((values or [])[:3]).lower()
            if not values:
                return True
            if any("tbd" in (v or "").lower() for v in values):
                return True
            if all(prefix.lower() in (v or "").lower() for v in values):
                return True
            return False

        async def _backfill_list(section: str, count: int, lines: int) -> list[str] | None:
            try:
                bf_prompt = f"""Create {count} concise items for the '{section}' section about this RFP. 
Each item must be at most {lines} lines, with each line ~18–28 words, domain-appropriate and specific.
Respond ONLY with a JSON array of strings of length {count}.
Context:
{rfp_text[:2000]}
"""
                text = await async_llm_complete(
                    messages=[
                        {"role": "system", "content": "Return valid JSON only. No prose."},
                        {"role": "user", "content": bf_prompt},
                    ],
                    temperature=0.4,
                    max_tokens=1200,
                )
                arr = json.loads(text)
                if isinstance(arr, list) and len(arr) >= count:
                    return [ _format_multiline(str(x), lines) for x in arr[:count] ]
            except Exception:
                return None
            return None

        async def _backfill_milestones(count: int) -> list[dict] | None:
            try:
                bf_prompt = f"""Create {count} milestones for this project with fields: phase, duration, description.
Description must be ≤4 lines (each ~16–24 words).
Respond ONLY with a JSON array of objects length {count}.
Context:
{rfp_text[:2000]}
"""
                text = await async_llm_complete(
                    messages=[
                        {"role": "system", "content": "Return valid JSON only. No prose."},
                        {"role": "user", "content": bf_prompt},
                    ],
                    temperature=0.35,
                    max_tokens=1200,
                )
                arr = json.loads(text)
                out = []
                if isinstance(arr, list):
                    for m in arr[:count]:
                        if isinstance(m, dict):
                            out.append({
                                "phase": str(m.get("phase") or "Phase").strip(),
                                "duration": str(m.get("duration") or "2 weeks").strip(),
                                "description": _format_multiline(m.get("description") or "", 4)
                            })
                if len(out) == count:
                    return out
            except Exception:
                return None
            return None

        # Backfill objectives/acceptance/tech stack/milestones when they look placeholder
        if key_fillers or _looks_placeholder_list(solution_data.get("key_challenges") or [], "challenge"):
            new_list = await _backfill_list("Key Challenges", 7, 5)
            if new_list:
                solution_data["key_challenges"] = _apply_backfill(new_list, 7, 5, "Challenge")
        if obj_fillers or _looks_placeholder_list(solution_data.get("objectives") or [], "objective"):
            new_list = await _backfill_list("Objectives", 7, 5)
            if new_list:
                solution_data["objectives"] = _apply_backfill(new_list, 7, 5, "Objective")
        if acc_fillers or _looks_placeholder_list(solution_data.get("acceptance_criteria") or [], "criterion"):
            new_list = await _backfill_list("Acceptance Criteria", 7, 4)
            if new_list:
                solution_data["acceptance_criteria"] = _apply_backfill(new_list, 7, 4, "Criterion")
        if stack_fillers or _looks_placeholder_list(solution_data.get("technical_stack") or [], "technology"):
            new_list = await _backfill_list("Technical Stack (technologies/tools/services)", 15, 1)
            if new_list:
                count = max(12, min(len(new_list), 20))
                solution_data["technical_stack"] = _apply_backfill(new_list, count, None, "Technology")
        if not solution_data.get("milestones") or any("tbd" in (m.get("description") or "").lower() for m in solution_data.get("milestones") or []):
            new_ms = await _backfill_milestones(7)
            if new_ms:
                solution_data["milestones"] = new_ms
        
        # Improve diagram if too basic
        diagram = solution_data.get('architecture_diagram')
        if diagram and _diagram_is_basic(diagram):
            safe_print("[INFO] Diagram is too basic, improving...")
            better = await asyncio.to_thread(_improve_diagram_mermaid, rfp_text, diagram)
            if better and better.strip():
                solution_data['architecture_diagram'] = _sanitize_mermaid_code(better)
        else:
            # Ensure diagram exists
            if not solution_data.get('architecture_diagram'):
                seed = "flowchart TD\nsubgraph CL[Client]\nWebApp[Web App]\nend\nsubgraph GW[Gateway]\nAPIGW[API Gateway]\nend\nsubgraph MS[Services]\nUserSvc[User Service]\nDataSvc[Data Service]\nAISvc[AI Service]\nend\nsubgraph DL[Data]\nDB[(Database)]\nCache[(Cache)]\nVectorDB[(Vector DB)]\nend\nWebApp --> APIGW\nAPIGW --> UserSvc\nAPIGW --> DataSvc\nAPIGW --> AISvc\nUserSvc --> DB\nDataSvc --> DB\nAISvc --> VectorDB\nUserSvc --> Cache"
                solution_data['architecture_diagram'] = _sanitize_mermaid_code(seed)
            else:
                solution_data['architecture_diagram'] = _sanitize_mermaid_code(solution_data.get('architecture_diagram'))
        
        # Render diagram to image
        if solution_data.get("architecture_diagram"):
            image_path = render_mermaid_to_image(solution_data["architecture_diagram"])
            if image_path and os.path.exists(image_path):
                try:
                    with open(image_path, "rb") as image_file:
                        encoded_image = base64.b64encode(image_file.read()).decode("utf-8")
                        solution_data["architecture_diagram_image"] = f"data:image/png;base64,{encoded_image}"
                    os.remove(image_path)
                except Exception as e:
                    safe_print(f"[WARN] Error reading generated diagram: {e}")
                    solution_data["architecture_diagram_image"] = None
            else:
                safe_print("[INFO] Diagram image generation failed; continuing without image.")
                solution_data["architecture_diagram_image"] = None
        
        return GeneratedSolution(**solution_data), retrieval_info
        
    except Exception as e:
        safe_print(f"Error with Groq API: {str(e)}")
        # Fallback solution (INR costs; includes years_of_experience)
        return GeneratedSolution(
            title="Technical Solution Proposal",
            date=datetime.now().strftime('%B %Y'),
            problem_statement="Based on the RFP requirements, we need to develop a comprehensive technical solution.",
            key_challenges=[
                "Integration complexity with existing systems",
                "Scalability and performance requirements",
                "Data security and compliance",
                "User adoption and training"
            ],
            solution_approach=[
                {
                    "title": "Step 1: Requirements Analysis & Architecture Design",
                    "description": "Comprehensive analysis of requirements and design of solution architecture"
                },
                {
                    "title": "Step 2: Technology Stack Selection & Setup",
                    "description": "Selection and configuration of appropriate technologies and frameworks"
                },
                {
                    "title": "Step 3: Development & Implementation",
                    "description": "Agile development approach with iterative implementation"
                },
                {
                    "title": "Step 4: Testing & Quality Assurance",
                    "description": "Comprehensive testing strategy including unit, integration, and user acceptance testing"
                },
                {
                    "title": "Step 5: Deployment & Go-Live",
                    "description": "Production deployment with monitoring and support setup"
                }
            ],
            architecture_diagram=_sanitize_mermaid_code("graph TD\nA[Client] --> B[API Gateway]\nB --> C[Microservice 1]\nB --> D[Microservice 2]"),
            milestones=[
                {"phase": "Planning & Design", "duration": "2 weeks", "description": "Requirements analysis and solution design"},
                {"phase": "Development Phase 1", "duration": "4 weeks", "description": "Core functionality development"},
                {"phase": "Development Phase 2", "duration": "4 weeks", "description": "Advanced features and integrations"},
                {"phase": "Testing & QA", "duration": "2 weeks", "description": "Comprehensive testing and bug fixes"},
                {"phase": "Deployment & Launch", "duration": "1 week", "description": "Production deployment and go-live"},
                {"phase": "Support & Maintenance", "duration": "Ongoing", "description": "Post-launch support and maintenance"}
            ],
            technical_stack=["Python", "React", "FastAPI", "PostgreSQL", "Docker", "AWS"],
            objectives=[
                "Deliver a scalable and robust technical solution",
                "Ensure seamless integration with existing systems",
                "Provide comprehensive documentation and training",
                "Achieve high user adoption and satisfaction"
            ],
            acceptance_criteria=[
                "System meets all functional requirements",
                "Performance benchmarks achieved",
                "Security compliance verified",
                "User acceptance testing completed successfully"
            ],
            key_performance_indicators=[
                {"metric": "System Uptime", "target": "99.9%", "measurement_method": "Automated monitoring tools", "frequency": "Continuous"},
                {"metric": "Response Time", "target": "< 2 seconds", "measurement_method": "Application performance monitoring", "frequency": "Hourly"},
                {"metric": "User Satisfaction", "target": "> 4.5/5", "measurement_method": "User surveys and feedback", "frequency": "Quarterly"},
                {"metric": "Bug Resolution Time", "target": "< 24 hours", "measurement_method": "Issue tracking system", "frequency": "Weekly"},
                {"metric": "System Security Score", "target": "> 90%", "measurement_method": "Security assessment tools", "frequency": "Monthly"}
            ],
            resources=[
                {"role": "Project Manager", "count": 1, "years_of_experience": 10, "responsibilities": "Project governance and stakeholder communication"},
                {"role": "Solution Architect", "count": 1, "years_of_experience": 8, "responsibilities": "Architecture and design oversight"},
                {"role": "Backend Engineer", "count": 2, "years_of_experience": 5, "responsibilities": "API and data layer development"},
                {"role": "Frontend Engineer", "count": 2, "years_of_experience": 4, "responsibilities": "UI implementation and integrations"}
            ],
            cost_analysis=[
                {"item": "Discovery & Design", "cost": "₹650,000", "notes": "Requirements and architecture"},
                {"item": "Implementation", "cost": "₹2,900,000", "notes": "Core features and integrations"},
                {"item": "Testing & QA", "cost": "₹580,000", "notes": "Automated and UAT"},
                {"item": "Deployment & Training", "cost": "₹420,000", "notes": "Go-live and enablement"}
            ]
        ), retrieval_info
# Document generation functions
def create_word_document(solution: GeneratedSolution) -> str:
    """Create a Word document from the generated solution"""
    doc = Document()

    from company_info import COMPANY_ABOUT

    # Initialize bookmark_id
    bookmark_id = 1

    # Make fields update on open
    try:
        settings = doc.settings
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.element.append(update_fields)
    except Exception:
        pass

    # Add page numbers in footer: Page X of Y
    try:
        _add_page_number_footer(doc)
    except Exception:
        pass

    # Cover
    logo_path = _get_logo_path()
    if os.path.exists(logo_path):
        try:
            picture = doc.add_picture(logo_path, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    title = doc.add_heading(solution.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(solution.date)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contents (Word TOC)
    doc.add_page_break()
    doc.add_heading('Contents', level=1)
    _insert_toc(doc)

    # About Company section
    doc.add_page_break()
    h = doc.add_heading('About Company', level=1)
    # Add bookmark without assigning the return value back
    _add_bookmark(h, 'sec_about_company', bookmark_id)
    
    # Split the company info into paragraphs and add them
    for paragraph_text in COMPANY_ABOUT.strip().split('\n\n'):
        if paragraph_text.strip():
            p = doc.add_paragraph()
            p.add_run(paragraph_text.strip())

    # Content sections
    doc.add_page_break()
    bookmark_id = 1

    def _flatten_text(text: str, join_newlines: bool = True) -> str:
        if text is None:
            return ""
        text = str(text)
        if join_newlines:
            return " ".join(text.split())
        return text.strip()

    def _add_paragraph_block(text: str, heading: bool = False) -> None:
        if not text:
            return
        blocks = re.split(r'\n\s*\n', str(text).strip())
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            para = doc.add_paragraph(_flatten_text(block))
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if not heading else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(6)

    def _set_cell_alignment(cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(6)

    h = doc.add_heading('Problem Statement', level=1)
    _add_bookmark(h, 'sec_problem_statement', bookmark_id)
    _add_paragraph_block(solution.problem_statement)

    h = doc.add_heading('Key Challenges', level=1)
    _add_bookmark(h, 'sec_key_challenges', bookmark_id)
    for challenge in solution.key_challenges:
        text = _flatten_text(challenge)
        if text:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
    h = doc.add_heading('Our Solution Approach', level=1)
    _add_bookmark(h, 'sec_solution_approach', bookmark_id)
    for i, step in enumerate(solution.solution_approach, 1):
        doc.add_heading(f'{step.title}', level=2)
        _add_paragraph_block(step.description)

    if solution.architecture_diagram or solution.architecture_diagram_image:
        h = doc.add_heading('Architecture Diagram', level=1)
        bookmark_id = _add_bookmark(h, 'sec_architecture_diagram', bookmark_id)
        image_added = False
        if solution.architecture_diagram_image:
            try:
                b64_data = solution.architecture_diagram_image.split(',', 1)[-1]
                image_bytes = base64.b64decode(b64_data)
                temp_path = os.path.join(tempfile.gettempdir(), f"arch_diagram_{datetime.now().timestamp()}.png")
                with open(temp_path, "wb") as img_file:
                    img_file.write(image_bytes)
                doc.add_picture(temp_path, width=Inches(6.5))
                os.remove(temp_path)
                image_added = True
            except Exception as ex:
                safe_print(f"[WARN] Failed to use base64 architecture diagram: {ex}")
        sanitized_diagram = _sanitize_mermaid_code(solution.architecture_diagram)
        if not image_added and sanitized_diagram:
            image_path = render_mermaid_to_image(sanitized_diagram)
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6.5))
                os.remove(image_path)
                image_added = True
            except Exception as e:
                safe_print(f"[WARN] Error adding diagram to document: {str(e)}")
            if image_path and os.path.exists(image_path):
                try:
                    os.remove(image_path)
                except Exception:
                    pass
        if not image_added and sanitized_diagram:
            safe_print("[INFO] Rendering to image failed; including Mermaid code as fallback.")
            doc.add_paragraph("Architecture Diagram (Mermaid code fallback):")
            doc.add_paragraph(sanitized_diagram)
    h = doc.add_heading('Technology Stack', level=1)
    _add_bookmark(h, 'sec_technical_stack', bookmark_id)
    if solution.technical_stack:
        cols = 3
        rows = math.ceil(len(solution.technical_stack) / cols)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for idx, tech in enumerate(solution.technical_stack):
            r = idx // cols
            c = idx % cols
            table.cell(r, c).text = _flatten_text(tech)
        for cell in table._cells:
            _set_cell_alignment(cell)
            for paragraph in cell.paragraphs:
                paragraph.style = 'Normal'

    h = doc.add_heading('Key Milestones', level=1)
    _add_bookmark(h, 'sec_key_milestones', bookmark_id)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Duration'
    hdr_cells[2].text = 'Description'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for milestone in solution.milestones:
        row_cells = table.add_row().cells
        row_cells[0].text = _flatten_text(milestone.phase)
        row_cells[1].text = _flatten_text(milestone.duration)
        row_cells[2].text = _flatten_text(milestone.description, join_newlines=False)
        for cell in row_cells:
            _set_cell_alignment(cell)

    h = doc.add_heading('Objectives', level=1)
    _add_bookmark(h, 'sec_objectives', bookmark_id)
    for objective in solution.objectives:
        text = _flatten_text(objective, join_newlines=False)
        if text:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)

    h = doc.add_heading('Acceptance Criteria', level=1)
    _add_bookmark(h, 'sec_acceptance_criteria', bookmark_id)
    for criteria in solution.acceptance_criteria:
        text = _flatten_text(criteria, join_newlines=False)
        if text:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)

    h = doc.add_heading('Resources', level=1)
    _add_bookmark(h, 'sec_resources', bookmark_id)
    r_table = doc.add_table(rows=1, cols=4)
    r_table.style = 'Table Grid'
    r_hdr = r_table.rows[0].cells
    r_hdr[0].text = 'Role'
    r_hdr[1].text = 'Count'
    r_hdr[2].text = 'Years of Experience'
    r_hdr[3].text = 'Responsibilities'
    for cell in r_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for res in solution.resources:
        row = r_table.add_row().cells
        row[0].text = _flatten_text(res.role)
        row[1].text = str(res.count)
        row[2].text = str(res.years_of_experience or '')
        row[3].text = _flatten_text(res.responsibilities or "", join_newlines=False)
        for cell in row:
            _set_cell_alignment(cell)

    h = doc.add_heading('Cost Analysis', level=1)
    _add_bookmark(h, 'sec_cost_analysis', bookmark_id)
    c_table = doc.add_table(rows=1, cols=3)
    c_table.style = 'Table Grid'
    c_hdr = c_table.rows[0].cells
    c_hdr[0].text = 'Item'
    c_hdr[1].text = 'Cost (INR)'
    c_hdr[2].text = 'Notes'
    for cell in c_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for ci in solution.cost_analysis:
        row = c_table.add_row().cells
        row[0].text = _flatten_text(ci.item)
        row[1].text = ci.cost
        row[2].text = _flatten_text(ci.notes or "", join_newlines=False)
        for cell in row:
            _set_cell_alignment(cell)

    # Add KPI Section
    h = doc.add_heading('Key Performance Indicators', level=1)
    _add_bookmark(h, 'sec_kpi', bookmark_id)
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = 'Table Grid'
    kpi_hdr = kpi_table.rows[0].cells
    kpi_hdr[0].text = 'Metric'
    kpi_hdr[1].text = 'Target'
    kpi_hdr[2].text = 'Measurement Method'
    kpi_hdr[3].text = 'Frequency'
    for cell in kpi_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for kpi in solution.key_performance_indicators:
        row = kpi_table.add_row().cells
        row[0].text = _flatten_text(kpi.metric)
        row[1].text = _flatten_text(kpi.target)
        row[2].text = _flatten_text(kpi.measurement_method, join_newlines=False)
        row[3].text = _flatten_text(kpi.frequency or "", join_newlines=False)
        for cell in row:
            _set_cell_alignment(cell)

    temp_dir = tempfile.gettempdir()
    doc_path = os.path.join(temp_dir, f'technical_proposal_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(doc_path)
    return doc_path

# API Endpoints
@app.post("/api/generate-solution", response_model=SolutionWithRecommendations)
async def generate_solution(file: UploadFile = File(...), method: str = "knowledgeBase", knowledge_base: Optional[str] = None):
    """Generate solution from uploaded RFP document"""
    logging.getLogger("sharepoint.flow").info("generate-solution called method=%s knowledge_base=%s", method, knowledge_base)
    
    # Validate file type (PDF and DOCX only)
    allowed_types = ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}. Only PDF and DOCX are supported.")
    
    # Create temporary file
    temp_dir = tempfile.gettempdir()
    temp_file_path = os.path.join(temp_dir, file.filename)
    
    try:
        # Save uploaded file while enforcing size limit
        with open(temp_file_path, "wb") as buffer:
            total_written = 0
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                total_written += len(chunk)
                if total_written > MAX_FILE_SIZE_BYTES:
                    raise HTTPException(status_code=400, detail=f"File too large. Max size is {MAX_FILE_SIZE_MB}MB.")
                buffer.write(chunk)
        await file.close()

        # Double-check file size on disk
        if os.path.getsize(temp_file_path) == 0:
            raise HTTPException(status_code=400, detail="Uploaded file is empty.")
        
        # Extract text based on file type
        if file.content_type == 'application/pdf':
            rfp_text = extract_text_from_pdf(temp_file_path)
        elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            rfp_text = extract_text_from_docx(temp_file_path)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")
        
        if not rfp_text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the document")
        
        # Generate solution using Groq
        if method == "llmOnly":
            solution, retrieval_info = await analyze_rfp_with_groq(rfp_text, use_rag=False)
        else:
            solution, retrieval_info = await analyze_rfp_with_groq(rfp_text, use_rag=True, knowledge_base=knowledge_base)
        
        recs = find_product_recommendations(solution.problem_statement, threshold=0.20)
        return SolutionWithRecommendations(solution=solution, recommendations=recs, retrieval_info=retrieval_info)
    
    except HTTPException:
        raise
    except Exception as e:
        safe_print(f"FATAL ERROR in /api/generate-solution: {e}") 
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    finally:
        # Cleanup temporary file
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception:
                pass

@app.post("/api/generate-solution-text", response_model=SolutionWithRecommendations)
async def generate_solution_text(body: GenerateTextBody):
    """Generate solution directly from a raw problem statement / use case text."""
    rfp_text = (body.text or "").strip()
    if not rfp_text:
        raise HTTPException(status_code=400, detail="Text is required")
    try:
        logging.getLogger("sharepoint.flow").info("generate-solution-text called method=%s knowledge_base=%s", body.method, body.knowledge_base)
        solution, retrieval_info = await analyze_rfp_with_groq(rfp_text, use_rag=(body.method != "llmOnly"), knowledge_base=body.knowledge_base)
        recs = find_product_recommendations(solution.problem_statement, threshold=0.20)
        return SolutionWithRecommendations(solution=solution, recommendations=recs, retrieval_info=retrieval_info)
    except Exception as e:
        safe_print(f"FATAL ERROR in /api/generate-solution-text: {e}") 
        raise HTTPException(status_code=500, detail=f"Error generating from text: {str(e)}")
@app.post("/api/recommendations", response_model=List[ProductRecommendation])
async def get_recommendations(body: RecommendBody):
    text = (body.text or "").strip()
    if not text:
        return []
    return find_product_recommendations(text, threshold=0.20)

@app.post("/api/solutions")
async def save_solution(solution: GeneratedSolution, x_user_email: Optional[str] = Header(None), db: Session = Depends(get_db)):
    """Save a generated solution to the database and filesystem"""
    try:
        # Create Word document and save to generated_solutions folder
        solutions_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generated_solutions')
        if not os.path.exists(solutions_dir):
            os.makedirs(solutions_dir)
            
        file_name = f"{solution.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc_path = os.path.join(solutions_dir, file_name)
        
        # Create the document
        doc = create_word_document(solution)
        shutil.move(doc, doc_path)
        
        # Save to database
        solution_record = DBSolution(
            title=solution.title,
            file_path=doc_path,
            user_id=(x_user_email or "anonymous")
        )
        db.add(solution_record)
        db.commit()
        db.refresh(solution_record)
        
        return {"id": solution_record.id}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving solution: {str(e)}")

@app.post("/api/download-solution")
async def download_solution(solution: GeneratedSolution):
    """Download generated solution as Word document"""
    
    try:
        # Create Word document
        doc_path = create_word_document(solution)
        
        # Return file response
        return FileResponse(
            doc_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename='technical_proposal.docx',
            headers={"Content-Disposition": "attachment; filename=technical_proposal.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating document: {str(e)}")

@app.get("/api/logo")
async def get_company_logo():
    """Serve the company logo image for frontend preview."""
    logo_path = _get_logo_path()
    if not os.path.exists(logo_path):
        raise HTTPException(status_code=404, detail="Logo not found")
    return FileResponse(logo_path, media_type='image/png', filename='AIONOS_logo.png')

@app.get("/api/solutions")
async def get_solutions(x_user_email: Optional[str] = Header(None), db: Session = Depends(get_db)):
    """Get list of generated solutions visible to the requester.
    Admin: sees only Admin-generated.
    Manager: sees Manager-generated and Admin-generated.
    Others: sees only their own if header provided; otherwise none.
    """
    requester = (x_user_email or "").strip()
    if not requester:
        solutions = []
    else:
        if requester.lower() == "manager@gmail.com":
            solutions = (
                db.query(DBSolution)
                .filter(DBSolution.user_id.in_(["Manager@gmail.com", "Admin@gmail.com"]))
                .order_by(DBSolution.generated_date.desc())
                .all()
            )
        else:
            # default and Admin path: only own
            solutions = (
                db.query(DBSolution)
                .filter(DBSolution.user_id == requester)
                .order_by(DBSolution.generated_date.desc())
                .all()
            )
    return [
        {
            "id": solution.id,
            "title": solution.title,
            "generated_date": solution.generated_date.isoformat(),
            "file_path": solution.file_path
        }
        for solution in solutions
    ]
 
@app.get("/api/solutions/{solution_id}")
async def get_solution(solution_id: int, x_user_email: Optional[str] = Header(None), db: Session = Depends(get_db)):
    """Get a specific solution by ID if requester has access."""
    solution = db.query(DBSolution).filter(DBSolution.id == solution_id).first()
    if not solution:
        raise HTTPException(status_code=404, detail="Solution not found")
    requester = (x_user_email or "").strip()
    if not requester:
        raise HTTPException(status_code=403, detail="Unauthorized")
    # Access control: Admin -> only Admin's; Manager -> Admin or Manager; Others -> only own
    allowed = False
    if requester.lower() == "manager@gmail.com":
        allowed = solution.user_id in ("Manager@gmail.com", "Admin@gmail.com")
    else:
        allowed = solution.user_id == requester
    if not allowed:
        raise HTTPException(status_code=403, detail="Forbidden")
   
    file_path = solution.file_path
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Solution file not found")
   
    return FileResponse(
        file_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=f'{solution.title}.docx'
    )

@app.post("/api/chat")
async def chat_with_groq(request: Request):
    """
    Basic Chatbot Agent for RFP App using Groq LLM.
    Handles navigation commands (jump to section) and Q&A from generated solution.
    """
    try:
        data = await request.json()
        message = (data.get("message") or "").strip()
        solution_title = data.get("solution_title")
        solution_content = data.get("solution_content")

        if not message:
            return {"response":"Please enter a message.","action": None}
        
        sections = {
            "problem statement": "problem-statement",
            "key challenges": "key-challenges",
            "solution approach": "solution-approach",
            "architecture diagram": "architecture-diagram",
            "milestones": "milestones",
            "technical stack": "technical-stack",
            "cost analysis": "cost-analysis",
            "objectives": "objectives",
            "acceptance criteria": "acceptance-criteria",
            "resources": "resources",
            "key performance indicators": "key-performance-indicators",
        }

        for key,sec_id in sections.items():
            if key in message.lower():
                if any(x in message.lower() for x in ["go to", "jump to", "show me", "take me to", "navigate to"]):
                    return {
                        "response": f"Navigating to {key} section.",
                        "action": {"type": "jump_to", "section": sec_id}
                    }
                
        # Check if solution content is available
        has_solution_content = bool(solution_content)
        
        # If no solution content, inform user they need to generate a solution first
        if not has_solution_content:
            # Only allow basic greetings when no solution is available
            greeting_keywords = ["hi", "hello", "hey", "greetings", "how are you", "what can you do"]
            if not any(kw in message.lower() for kw in greeting_keywords):
                return {
                    "response": "I can only answer questions about your generated RFP solution. Please upload an RFP and generate a solution first, then ask me about its content.",
                    "action": None
                }
        
        # Parse and format solution content for better LLM understanding
        solution_context = ""
        if solution_content:
            try:
                # Parse JSON string if it's a string
                if isinstance(solution_content, str):
                    solution_data = json.loads(solution_content)
                else:
                    solution_data = solution_content
                
                # Format the solution data in a readable way for the LLM
                tech_stack = solution_data.get('technical_stack', [])
                tech_stack_str = ', '.join(tech_stack) if tech_stack and len(tech_stack) > 0 else 'Not specified'
                
                # Format solution approach safely
                solution_approach_list = solution_data.get('solution_approach', [])
                approach_lines = []
                for i, step in enumerate(solution_approach_list[:3]):
                    if isinstance(step, dict):
                        title = step.get('title', 'Step')
                        desc = step.get('description', '')[:100]
                        approach_lines.append(f"  {i+1}. {title}: {desc}...")
                    else:
                        approach_lines.append(f"  {i+1}. {str(step)[:100]}...")
                
                # Format milestones safely
                milestones_list = solution_data.get('milestones', [])
                milestone_lines = []
                for i, milestone in enumerate(milestones_list[:3]):
                    if isinstance(milestone, dict):
                        phase = milestone.get('phase', 'Phase')
                        duration = milestone.get('duration', 'N/A')
                        milestone_lines.append(f"  {i+1}. {phase}: {duration}")
                    else:
                        milestone_lines.append(f"  {i+1}. {str(milestone)[:100]}...")
                
                # Format cost analysis safely
                cost_analysis_list = solution_data.get('cost_analysis', [])
                cost_lines = []

                # Identify any explicit total row (e.g., item name contains "total")
                explicit_total_value = None
                explicit_total_label = None
                total_entries = len(cost_analysis_list)

                for cost_item in cost_analysis_list:
                    if isinstance(cost_item, dict):
                        item_name = (cost_item.get('item') or "").lower()
                        if "total" in item_name and explicit_total_value is None:
                            explicit_total_value = cost_item.get('cost')
                            explicit_total_label = cost_item.get('item')

                # Prepare preview lines for the first few items
                missing_cost_values = 0
                for i, cost_item in enumerate(cost_analysis_list[:5]):
                    if isinstance(cost_item, dict):
                        item = cost_item.get('item', f'Item {i+1}')
                        cost = cost_item.get('cost', 'N/A')
                        notes = cost_item.get('notes', '')
                        if notes:
                            cost_lines.append(f"  {i+1}. {item}: {cost} ({notes[:120]}...)")
                        else:
                            cost_lines.append(f"  {i+1}. {item}: {cost}")
                        if not cost or cost.strip().lower() in ("n/a", "na", "not available"):
                            missing_cost_values += 1
                    else:
                        cost_lines.append(f"  {i+1}. {str(cost_item)[:120]}...")
                        missing_cost_values += 1
                if len(cost_analysis_list) > 5:
                    cost_lines.append(f"  ...and {total_entries - 5} additional cost items.")

                # Calculate total cost summary
                if explicit_total_value:
                    total_cost_summary = f"  Total Cost (as provided): {explicit_total_value}"
                elif total_entries > 0:
                    total_cost_summary = "  Total Cost: Not explicitly provided; list shows individual cost items."
                else:
                    total_cost_summary = "  No cost breakdown provided."
                
                solution_context = f"""
Solution Title: {solution_data.get('title', solution_title or 'N/A')}
Date: {solution_data.get('date', 'N/A')}

Problem Statement: {solution_data.get('problem_statement', 'N/A')[:500]}...

Key Challenges: {', '.join([str(c) for c in solution_data.get('key_challenges', [])[:5]])}...

Technical Stack: {tech_stack_str}

Solution Approach: {len(solution_approach_list)} steps defined
{chr(10).join(approach_lines) if approach_lines else '  No steps defined'}

Objectives: {len(solution_data.get('objectives', []))} objectives defined
- {chr(10).join([f"  • {str(obj)[:100]}..." for obj in solution_data.get('objectives', [])[:3]]) if solution_data.get('objectives') else '  No objectives defined'}

Milestones: {len(milestones_list)} phases
{chr(10).join(milestone_lines) if milestone_lines else '  No milestones defined'}

Resources: {len(solution_data.get('resources', []))} roles defined

Cost Analysis: {len(solution_data.get('cost_analysis', []))} items
{chr(10).join(cost_lines) if cost_lines else '  No cost breakdown provided'}
{total_cost_summary}

Key Performance Indicators: {len(solution_data.get('key_performance_indicators', []))} KPIs defined
"""
            except (json.JSONDecodeError, TypeError, AttributeError) as e:
                safe_print(f"[WARN] Failed to parse solution content: {e}")
                solution_context = f"Solution Content (raw): {solution_content[:1000]}"
        else:
            solution_context = "No solution content available."
                
        prompt = f"""
You are a friendly AI assistant integrated into an RFP Solution Generator app.
Your role is to answer questions about the generated RFP solution based on the context provided below.

CRITICAL RULES:
1. You MUST answer questions based ONLY on the solution content provided below.
2. You MUST NOT use any external knowledge or general information.
3. If specific details are missing, clearly state: "This information is not available in the generated RFP solution," and explain any assumptions or gaps.
4. Only use the redirect message ("I can only answer questions about your generated RFP solution...") if the question is clearly unrelated (e.g., weather, sports, jokes). Otherwise, do your best to answer using the solution context.
5. When summarising or listing items, format each bullet or numbered point on its own line using plain text (e.g., "1. Item detail"). Avoid duplicate asterisks, extra Markdown, or multiple points on the same line.
6. If cost amounts are provided, calculate totals or subtotals directly from those figures. If any amounts are missing, explain which items lack data and whether the total is partial or unavailable.
7. Be accurate, helpful, and concise. When comparing or aligning items (e.g., objectives vs. use case), reference the relevant parts of the solution and describe how they relate.

Solution Content:
{solution_context}

User Question: {message}

Provide a helpful answer based ONLY on the solution content above. If the question cannot be fully answered from the solution content, explain what is available and what information is missing.
"""
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You are a helpful assistant for technical RFP proposal app. Answer questions accurately based on the provided solution content."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=500
        )
        answer = response.choices[0].message.content.strip()
        answer = _format_list_markers(answer)
        return {"response": answer, "action": None}
    except Exception as e:
        safe_print("Chat API Error:", str(e))
        import traceback
        traceback.print_exc()
        return {"response": "Sorry, something went wrong while generating a reply.", "action": None}

@app.post("/api/tender-chat")
async def chat_with_tenders(request: Request):
    """
    Tender-specific Chatbot Agent using Groq LLM.
    Handles questions about scraped tender data with real-time access.
    """
    try:
        data = await request.json()
        message = (data.get("message") or "").strip()
        tender_data = data.get("tender_data", [])

        if not message:
            return {"response": "Please enter a message.", "action": None}
        
        # Get fresh tender data if not provided
        if not tender_data:
            try:
                safe_print("Fetching fresh tender data...")
                from scraper_service import fetch_all_sources
                fresh_data = fetch_all_sources(limit_per_source=50, ttlh_only=True, max_pages=1) or {}
                safe_print(f"Fresh data sources: {list(fresh_data.keys()) if fresh_data else 'None'}")
                tender_data = []
                for source_name, source_tenders in fresh_data.items():
                    safe_print(f"Processing {source_name}: {len(source_tenders)} tenders")
                    for tender in source_tenders:
                        tender_data.append({
                            "source": source_name,
                            "tender_id": tender.get("tender_id"),
                            "title": tender.get("title"),
                            "organization": tender.get("organization"),
                            "sector": tender.get("sector"),
                            "deadline": tender.get("deadline"),
                            "value": tender.get("value"),
                            "url": tender.get("url"),
                            "description": tender.get("description", "")
                        })
                safe_print(f"Total tender data prepared: {len(tender_data)} tenders")
            except Exception as e:
                safe_print(f"Error fetching fresh tender data: {e}")
                import traceback
                traceback.print_exc()
                tender_data = []

        # Create context from tender data
        tender_context = ""
        if tender_data:
            tender_context = f"Current Tender Data ({len(tender_data)} tenders):\n"
            for i, tender in enumerate(tender_data[:20], 1):  # Limit to first 20 tenders for context
                tender_context += f"""
{i}. Tender ID: {tender.get('tender_id', 'N/A')}
   Title: {tender.get('title', 'N/A')}
   Organization: {tender.get('organization', 'N/A')}
   Sector: {tender.get('sector', 'N/A')}
   Deadline: {tender.get('deadline', 'N/A')}
   Value: {tender.get('value', 'N/A')}
   Source: {tender.get('source', 'N/A')}
   Description: {tender.get('description', 'N/A')[:200]}...
   URL: {tender.get('url', 'N/A')}
"""
        else:
            tender_context = "No tender data available at the moment."

        prompt = f"""
        You are a helpful AI assistant for a tender management system. You have access to real-time scraped tender data and can answer questions about specific tenders, sectors, deadlines, values, and other details.

        {tender_context}

        User Question: {message}

        Instructions:
        - Answer questions about specific tenders by referencing their tender ID, title, organization, sector, deadline, value, etc.
        - Help users find tenders by sector, organization, deadline, or value range
        - Provide specific details when asked about particular tenders
        - If asked about a tender by number, refer to the tender ID and provide key details
        - Be helpful and specific in your responses
        - If you don't have information about a specific tender, say so clearly
        - Use the tender data provided above to give accurate answers
        """

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You are a helpful assistant for tender management and analysis. You have access to real-time tender data and can answer questions about specific tenders, sectors, deadlines, values, and other details."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )
        answer = response.choices[0].message.content.strip()
        return {"response": answer, "action": None}
    except Exception as e:
        safe_print("Tender Chat API Error:", str(e))
        return {"response": "Sorry, something went wrong while generating a reply. Please try again.", "action": None}
    
@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)